"""
赛事附件下载与处理服务
"""
import os
import uuid
import mimetypes
from typing import Optional, List, Dict, Any
from pathlib import Path
from datetime import datetime

import httpx
from sqlalchemy.orm import Session

from ...models import Contest, ContestAttachment
from ...core.config import get_settings
from .captcha_downloader import get_captcha_downloader

settings = get_settings()


class AttachmentService:
    """附件处理服务"""

    def __init__(self, db: Session):
        self.db = db
        self.upload_dir = Path(settings.UPLOAD_DIR) / "contest_attachments"
        self.upload_dir.mkdir(parents=True, exist_ok=True)

    def download_attachment(
        self,
        url: str,
        name: str,
        contest_id: str,
    ) -> Optional[ContestAttachment]:
        """下载附件"""
        # 先检查是否已存在
        existing = self.db.query(ContestAttachment).filter(
            ContestAttachment.contest_id == contest_id,
            ContestAttachment.url == url
        ).first()

        if existing:
            return existing

        # 下载文件
        try:
            # 尝试用验证码下载器
            downloader = get_captcha_downloader()
            file_content = downloader.download_with_captcha(url, max_attempts=2)

            if not file_content:
                # 验证码下载失败，尝试直接下载
                print("验证码下载失败，尝试直接下载...")
                headers = {
                    "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36"
                }
                response = httpx.get(url, headers=headers, timeout=60, follow_redirects=True)
                response.raise_for_status()
                file_content = response.content

            # 保存文件
            # 构造一个假的headers用于guess_extension
            fake_headers = {}
            ext = self._guess_extension(url, name, fake_headers)
            file_id = str(uuid.uuid4())
            filename = f"{file_id}{ext}"
            file_path = self.upload_dir / filename

            with open(file_path, "wb") as f:
                f.write(file_content)

            # 创建数据库记录
            attachment = ContestAttachment(
                id=file_id,
                contest_id=contest_id,
                name=name,
                url=url,
                file_path=str(file_path),
                file_type=ext.lstrip(".").lower() if ext else None,
                file_size=len(file_content),
                is_downloaded=True,
                is_parsed=False,
            )

            self.db.add(attachment)
            self.db.commit()
            self.db.refresh(attachment)

            return attachment

        except Exception as e:
            print(f"下载附件失败 {url}: {e}")
            # 即使下载失败，也创建一个记录
            attachment = ContestAttachment(
                id=str(uuid.uuid4()),
                contest_id=contest_id,
                name=name,
                url=url,
                is_downloaded=False,
                is_parsed=False,
            )
            self.db.add(attachment)
            self.db.commit()
            self.db.refresh(attachment)
            return attachment

    def _guess_extension(self, url: str, name: str, headers: dict) -> str:
        """猜测文件扩展名"""
        # 1. 从文件名猜测
        for ext in [".pdf", ".doc", ".docx", ".xls", ".xlsx", ".ppt", ".pptx", ".zip", ".rar"]:
            if ext.lower() in name.lower() or ext.lower() in url.lower():
                return ext

        # 2. 从 Content-Type 猜测
        content_type = headers.get("content-type", "")
        ext = mimetypes.guess_extension(content_type)
        if ext:
            return ext

        # 3. 默认
        return ".bin"

    def parse_attachment(self, attachment: ContestAttachment) -> Optional[str]:
        """解析附件内容"""
        if not attachment.is_downloaded or not attachment.file_path:
            return None

        file_path = Path(attachment.file_path)
        if not file_path.exists():
            return None

        ext = file_path.suffix.lower()
        content = None

        try:
            if ext == ".pdf":
                content = self._parse_pdf(file_path)
            elif ext in [".txt", ".md", ".text"]:
                content = self._parse_text(file_path)
            elif ext in [".doc", ".docx"]:
                content = self._parse_word(file_path)
            # 其他格式暂不处理

            if content:
                attachment.parsed_content = content
                attachment.is_parsed = True
                self.db.commit()

            return content

        except Exception as e:
            print(f"解析附件失败 {file_path}: {e}")
            return None

    def _parse_pdf(self, file_path: Path) -> Optional[str]:
        """解析PDF"""
        try:
            import PyPDF2
            text_parts = []
            with open(file_path, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
            return "\n\n".join(text_parts) if text_parts else None
        except ImportError:
            print("PyPDF2 未安装")
            return None
        except Exception as e:
            print(f"PDF解析失败: {e}")
            return None

    def _parse_text(self, file_path: Path) -> Optional[str]:
        """解析文本文件"""
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                return f.read()
        except UnicodeDecodeError:
            try:
                with open(file_path, "r", encoding="gbk") as f:
                    return f.read()
            except Exception:
                return None

    def _parse_word(self, file_path: Path) -> Optional[str]:
        """解析Word文档"""
        try:
            from docx import Document
            doc = Document(file_path)
            text_parts = []

            # 提取段落
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)

            # 提取表格
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                    if row_text:
                        text_parts.append(row_text)

            return "\n\n".join(text_parts) if text_parts else None
        except ImportError:
            print("python-docx 未安装")
            return None
        except Exception as e:
            print(f"Word解析失败: {e}")
            return None

    def get_attachments_by_contest(self, contest_id: str) -> List[ContestAttachment]:
        """获取赛事的所有附件"""
        return self.db.query(ContestAttachment).filter(
            ContestAttachment.contest_id == contest_id
        ).all()

    def get_attachment_content(self, attachment: ContestAttachment) -> Optional[str]:
        """获取附件解析后的内容，如果未解析则尝试解析"""
        if attachment.is_parsed and attachment.parsed_content:
            return attachment.parsed_content

        if attachment.is_downloaded:
            return self.parse_attachment(attachment)

        return None


def create_contest_attachments(
    db: Session,
    contest: Contest,
    attachments_data: List[Dict[str, str]],
    download: bool = True,
) -> List[ContestAttachment]:
    """创建赛事附件记录"""
    attachments = []
    service = AttachmentService(db)

    for att_data in attachments_data:
        name = att_data.get("name", "附件")
        url = att_data.get("url")

        if not url:
            continue

        # 过滤无效URL（比如本地文件路径 file:///...）
        url_lower = url.lower()
        if url_lower.startswith("file://") or url_lower.startswith("file:"):
            print(f"跳过无效URL（本地文件）: {url[:100]}...")
            continue

        # 只处理http/https链接
        if not url_lower.startswith("http://") and not url_lower.startswith("https://"):
            print(f"跳过非HTTP(S)链接: {url[:100]}...")
            continue

        if download:
            attachment = service.download_attachment(url, name, contest.id)
        else:
            # 只创建记录，不下载
            existing = db.query(ContestAttachment).filter(
                ContestAttachment.contest_id == contest.id,
                ContestAttachment.url == url
            ).first()

            if existing:
                attachments.append(existing)
                continue

            attachment = ContestAttachment(
                id=str(uuid.uuid4()),
                contest_id=contest.id,
                name=name,
                url=url,
                is_downloaded=False,
                is_parsed=False,
            )
            db.add(attachment)
            db.commit()
            db.refresh(attachment)

        attachments.append(attachment)

    return attachments
